from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"C:\ai\AI_for_Professional_24Days_Adarsha_Refined.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F7FA"
WHITE = "FFFFFF"
GRAY = "666666"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=60, start=120, bottom=60, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_cell_text(cell, size=9, bold=False, color="000000", align=WD_ALIGN_PARAGRAPH.LEFT):
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.04
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = RGBColor.from_string(color)


def repeat_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_schedule_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = [720, 2160, 4320, 2160]
    set_table_geometry(table, widths)
    headers = ["Day", "Topic", "What participants will practise and create", "Main AI tools"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, BLUE)
        style_cell_text(cell, size=9, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    repeat_header_row(table.rows[0])

    for day, topic, practice, tools in rows:
        cells = table.add_row().cells
        values = [day, topic, practice, tools]
        for idx, value in enumerate(values):
            cells[idx].text = value
            if idx == 0:
                set_cell_shading(cells[idx], PALE_BLUE)
                style_cell_text(cells[idx], size=9, bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
            elif idx == 1:
                style_cell_text(cells[idx], size=9, bold=True, color=DARK_BLUE)
            else:
                style_cell_text(cells[idx], size=9)
    doc.add_paragraph()


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for style_name, size, color, before, after in [
    ("Title", 24, DARK_BLUE, 0, 5),
    ("Subtitle", 11, GRAY, 0, 14),
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11, DARK_BLUE, 8, 4),
]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = style_name != "Subtitle"
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

list_style = styles["List Bullet"]
list_style.font.name = "Calibri"
list_style.font.size = Pt(10.5)
list_style.paragraph_format.left_indent = Inches(0.375)
list_style.paragraph_format.first_line_indent = Inches(-0.188)
list_style.paragraph_format.space_after = Pt(4)
list_style.paragraph_format.line_spacing = 1.15

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("AI for Professionals")
subtitle = doc.add_paragraph(style="Subtitle")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("24-Day Practical Training Programme | Adarsha Secondary School, Biratnagar")

lead = doc.add_paragraph()
lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = lead.add_run("A simple, hands-on programme for teachers, school staff, working professionals, and beginners.")
run.bold = True
run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

doc.add_heading("Programme Overview", level=1)
overview = doc.add_table(rows=5, cols=2)
overview.style = "Table Grid"
set_table_geometry(overview, [2700, 6660])
overview_rows = [
    ("Duration", "24 days, including 20 teaching days, two presentation days, assessment, and closing"),
    ("Session Length", "4 hours per day, including break"),
    ("Training Style", "Beginner-friendly, demonstration-based, and focused on useful results"),
    ("Group Project", "Prepared during Days 6-10 and presented on Day 11"),
    ("Individual Project", "Built during Days 17-20 and presented on Day 21"),
]
for idx, (label, detail) in enumerate(overview_rows):
    overview.rows[idx].cells[0].text = label
    overview.rows[idx].cells[1].text = detail
    set_cell_shading(overview.rows[idx].cells[0], LIGHT_BLUE)
    style_cell_text(overview.rows[idx].cells[0], bold=True, color=DARK_BLUE)
    style_cell_text(overview.rows[idx].cells[1])

doc.add_heading("Core AI Tools", level=1)
doc.add_paragraph(
    "The programme keeps the tool list focused so participants can learn deeply without feeling overwhelmed."
)
for text in [
    "ChatGPT, Claude, and Gemini for writing, planning, ideas, teaching materials, and project support.",
    "Perplexity for source-based research and working with multiple uploaded documents.",
    "NotebookLM for learning from trusted sources, study guides, summaries, and document questions.",
    "Google Veo 3 and Flow for simple AI video creation and visual storytelling.",
]:
    add_bullet(doc, text)

doc.add_heading("24-Day Syllabus", level=1)

doc.add_heading("Week 1: AI Foundations and Everyday Use", level=2)
add_schedule_table(doc, [
    ("1", "Introduction to AI", "Understand what AI can and cannot do. Compare answers from different AI assistants and list useful tasks for work or school.", "ChatGPT, Claude, Gemini"),
    ("2", "How to Give Good Instructions", "Turn unclear requests into clear prompts. Build a small personal prompt library for daily use.", "ChatGPT, Claude, Gemini"),
    ("3", "Choosing the Right AI Assistant", "Compare the strengths of ChatGPT, Claude, and Gemini and choose the best one for common tasks.", "ChatGPT, Claude, Gemini"),
    ("4", "AI for Daily Productivity", "Create a weekly plan, simplify a difficult message, plan an event or trip, and make a useful checklist.", "ChatGPT, Gemini"),
    ("5", "Safe and Responsible AI Use", "Learn about incorrect answers, privacy, bias, copyright, and fact-checking. Create a simple AI do-and-don't guide.", "ChatGPT, Gemini, Perplexity"),
])

doc.add_heading("Week 2: Writing, Research, and Teaching Resources", level=2)
add_schedule_table(doc, [
    ("6", "Professional Writing with AI", "Draft and improve emails, letters, reports, notices, and proposals. Practise formal, friendly, and concise writing.", "ChatGPT, Claude, Gemini"),
    ("7", "Perplexity: Learn from Multiple Documents", "Upload several documents, ask questions across them, compare information, find key points, and create a source-based summary.", "Perplexity"),
    ("8", "NotebookLM for Learning and Research", "Upload trusted sources and create summaries, FAQs, study guides, timelines, and an audio overview.", "NotebookLM"),
    ("9", "Flashcards and Interactive Quizzes", "Turn a chapter or document into flashcards and quizzes. Ask AI to create a simple HTML quiz, open it in a browser, and improve its design.", "ChatGPT, Claude, Gemini"),
    ("10", "AI for Teachers and Classrooms", "Create a lesson plan, worksheet, marking rubric, classroom activity, and age-appropriate explanation.", "ChatGPT, Claude, Gemini, NotebookLM"),
])

doc.add_heading("Day 11: Group Project Presentations", level=2)
doc.add_paragraph(
    "Teams present a practical AI-assisted resource such as a school resource pack, awareness campaign, research brief, or community information kit. Each team explains the tools used, shows the final output, and receives simple peer feedback."
)

doc.add_heading("Week 3: Presentations, Visuals, Video, and Useful Content", level=2)
add_schedule_table(doc, [
    ("12", "Presentations with AI", "Plan a clear presentation, create slide content, improve titles, and prepare simple speaker notes.", "ChatGPT, Claude, Gemini"),
    ("13", "Images, Posters, and Classroom Visuals", "Write strong image prompts and create a poster, classroom visual, or professional graphic.", "ChatGPT, Gemini"),
    ("14", "AI Video with Veo 3 and Flow", "Write a short script, plan scenes, generate video clips, and arrange them into a simple story.", "Google Veo 3, Flow, Gemini"),
    ("15", "Social and School Communication", "Create a short content plan, announcements, captions, and a small awareness campaign for a school or organisation.", "ChatGPT, Claude, Gemini"),
    ("16", "Understanding Data with AI", "Use sample marks, attendance, survey, or expense data to find patterns, write a summary, and suggest a simple chart.", "ChatGPT, Claude, Gemini"),
])

doc.add_heading("Week 4: Practical Career and Teacher Projects", level=2)
add_schedule_table(doc, [
    ("17", "Build a Professional Portfolio with AI", "Plan and create a simple portfolio that shows skills, teaching resources, achievements, and selected work.", "ChatGPT, Claude, Gemini"),
    ("18", "CV, Cover Letter, and Professional Profile", "Improve a CV, write a tailored cover letter, prepare a short professional biography, and practise interview questions.", "ChatGPT, Claude, Gemini"),
    ("19", "Build a Useful Teacher or Professional Project", "Choose and build one practical project such as a teacher resource pack, subject guide, interactive quiz, research brief, or school communication kit.", "Selected core AI tools"),
    ("20", "Project Polish and Peer Review", "Check accuracy, improve design and writing, collect peer feedback, and prepare a short presentation of the finished project.", "Selected core AI tools"),
])

doc.add_heading("Final Days", level=2)
add_schedule_table(doc, [
    ("21", "Individual Project Presentations", "Participants present what they built, how AI helped, and what they checked or improved themselves.", "Selected core AI tools"),
    ("22", "Review and Assessment Preparation", "Review the main tools and skills, practise common tasks, and ask final questions.", "All core AI tools"),
    ("23", "Final Assessment", "Complete a short written check, one practical AI task, and a brief reflection on future use.", "Selected core AI tools"),
    ("24", "Closing and Certificates", "Share selected projects, celebrate progress, discuss next steps, and distribute certificates.", "No new tools"),
])

doc.add_heading("Project Options for Days 17-20", level=1)
for text in [
    "A professional portfolio with a short biography, achievements, and selected work.",
    "An improved CV, cover letter, professional profile, and interview practice pack.",
    "A teacher resource pack with lesson plan, worksheet, quiz, rubric, and classroom activity.",
    "An interactive HTML quiz or flashcard set for students.",
    "A research or subject guide built from trusted documents using Perplexity or NotebookLM.",
    "A school communication kit with notices, presentation content, visuals, and a short video.",
]:
    add_bullet(doc, text)

doc.add_heading("Simple Assessment Plan", level=1)
assessment = doc.add_table(rows=1, cols=3)
assessment.style = "Table Grid"
set_table_geometry(assessment, [2700, 1800, 4860])
for idx, text in enumerate(["Part", "Weight", "What is checked"]):
    assessment.rows[0].cells[idx].text = text
    set_cell_shading(assessment.rows[0].cells[idx], BLUE)
    style_cell_text(assessment.rows[0].cells[idx], bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
repeat_header_row(assessment.rows[0])
for part, weight, check in [
    ("Daily practical work", "30%", "Participation and useful outputs from teaching days"),
    ("Group presentation", "20%", "Teamwork, clarity, and practical use of AI"),
    ("Individual project", "30%", "Usefulness, quality, accuracy, and responsible AI use"),
    ("Final assessment", "20%", "Understanding and ability to complete a practical task"),
]:
    cells = assessment.add_row().cells
    for idx, value in enumerate([part, weight, check]):
        cells[idx].text = value
        style_cell_text(cells[idx], align=WD_ALIGN_PARAGRAPH.CENTER if idx == 1 else WD_ALIGN_PARAGRAPH.LEFT)

doc.add_paragraph()
closing = doc.add_table(rows=1, cols=1)
closing.style = "Table Grid"
set_table_geometry(closing, [9360])
closing.rows[0].cells[0].text = (
    "This training is not about becoming an AI Engineer. "
    "It is about becoming an AI-powered teacher, professional, and creator."
)
set_cell_shading(closing.rows[0].cells[0], LIGHT_BLUE)
style_cell_text(closing.rows[0].cells[0], size=11, bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run("AI for Professionals | Adarsha Secondary School")
footer_run.font.name = "Calibri"
footer_run.font.size = Pt(8)
footer_run.font.color.rgb = RGBColor.from_string(GRAY)

doc.save(OUTPUT)
print(OUTPUT)
